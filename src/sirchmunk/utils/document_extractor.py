# Copyright (c) ModelScope Contributors. All rights reserved.
"""Unified document extraction facade over kreuzberg.

Centralizes all kreuzberg interaction into a single module, providing a clean,
configurable interface for document text extraction with support for tables,
metadata, language detection, OCR, and page-range filtering.

All other modules should import from here rather than from kreuzberg directly.
"""

from __future__ import annotations

import asyncio
import dataclasses
import multiprocessing as mp
import os
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, ClassVar, List, Optional, Sequence, Union
import xml.etree.ElementTree as ET

from loguru import logger


# ---------------------------------------------------------------------------
# Subprocess extraction helpers (module-level for picklability)
# ---------------------------------------------------------------------------

_EXTRACT_TIMEOUT_S = 600


def _extraction_worker(
    file_path: str,
    profile_dict: dict[str, Any],
    pipe_w: mp.connection.Connection,
) -> None:
    """Child process entry point: run kreuzberg, send result via pipe, exit.

    Sends a plain dict so no native kreuzberg/Rust objects cross the
    process boundary.  On failure sends ``{"_error": "<message>"}``.
    """
    try:
        import asyncio as _aio

        async def _run() -> dict[str, Any]:
            from sirchmunk.utils.document_extractor import (
                DocumentExtractor,
                ExtractionProfile,
            )
            profile = ExtractionProfile(**profile_dict)
            output = await DocumentExtractor.extract(file_path, profile)
            return {
                "content": output.content,
                "mime_type": output.mime_type,
                "metadata": output.metadata,
                "tables": output.tables,
                "detected_languages": output.detected_languages,
                "page_count": output.page_count,
            }

        pipe_w.send(_aio.run(_run()))
    except BaseException as exc:
        try:
            pipe_w.send({"_error": str(exc)})
        except Exception:
            pass
    finally:
        pipe_w.close()


def _run_extraction_in_child(
    file_path: str,
    profile_dict: dict[str, Any],
) -> dict[str, Any]:
    """Spawn an isolated child process, wait for its result.

    Unlike ``ProcessPoolExecutor``, a crash in one child never
    poisons future extractions — each call spawns a fresh process.
    """
    pipe_r, pipe_w = mp.Pipe(duplex=False)
    proc = mp.Process(
        target=_extraction_worker,
        args=(file_path, profile_dict, pipe_w),
        daemon=True,
    )
    proc.start()
    pipe_w.close()

    try:
        if not pipe_r.poll(timeout=_EXTRACT_TIMEOUT_S):
            proc.kill()
            proc.join(timeout=10)
            raise RuntimeError(
                f"Extraction timed out after {_EXTRACT_TIMEOUT_S}s"
            )
        result = pipe_r.recv()
    except EOFError:
        proc.join(timeout=10)
        raise RuntimeError(
            f"Worker crashed (exit code {proc.exitcode})"
        )
    finally:
        pipe_r.close()

    proc.join(timeout=30)
    if proc.is_alive():
        proc.kill()
        proc.join()

    if isinstance(result, dict) and "_error" in result:
        raise RuntimeError(result["_error"])
    return result


# ---------------------------------------------------------------------------
# Configuration profile
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class ExtractionProfile:
    """Immutable extraction configuration profile.

    Controls which kreuzberg features are enabled during document extraction.
    Default values align with the legacy ``fast_extract()`` behavior
    (plain text only, no extras).
    """

    output_format: str = "plain"
    """Output format: ``plain`` | ``markdown`` | ``html`` | ``djot``."""

    extract_tables: bool = False
    """Whether to extract and return tables."""

    extract_metadata: bool = False
    """Whether to return document metadata."""

    detect_language: bool = False
    """Whether to detect document language."""

    ocr_enabled: bool = False
    """Whether to enable OCR fallback."""

    ocr_backend: str = "tesseract"
    """OCR engine: ``tesseract`` | ``easyocr`` | ``paddleocr``."""

    ocr_language: str = "eng"
    """OCR language code (e.g. ``eng``, ``chi_sim``)."""

    page_start: Optional[int] = None
    """Page range start (0-indexed). ``None`` means first page."""

    page_end: Optional[int] = None
    """Page range end (inclusive). ``None`` means last page."""

    pdf_extract_images: bool = False
    """Extract images embedded in PDF pages."""

    pdf_extract_metadata: bool = False
    """Extract PDF-level metadata (author, title, etc.)."""

    force_ocr: bool = False
    """Force OCR for all pages, bypassing native text extraction.

    Maps directly to kreuzberg's ``ExtractionConfig.force_ocr``.
    Note: kreuzberg does not offer a "fallback" OCR mode —
    when set, OCR is always applied regardless of text layer presence.
    """

    force_ocr_pages: Optional[tuple[int, ...]] = None
    """Force OCR on specific pages only (0-indexed).

    Maps to kreuzberg's ``ExtractionConfig.force_ocr_pages``.
    Mutually exclusive with :attr:`force_ocr` — when both are set,
    ``force_ocr`` takes precedence.
    """

    pdf_password: Optional[str] = None
    """Password for encrypted PDFs."""

    max_concurrent: Optional[int] = None
    """Max concurrency for batch extraction."""


# ---------------------------------------------------------------------------
# Extraction output
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class ExtractionOutput:
    """Structured extraction result.

    Always contains ``content``.  Other fields are populated based on the
    :class:`ExtractionProfile` settings used during extraction.
    """

    content: str
    """Extracted text content."""

    mime_type: str = ""
    """MIME type of the source document."""

    metadata: dict[str, Any] = field(default_factory=dict)
    """Document metadata (empty when ``extract_metadata`` is disabled)."""

    tables: list[dict[str, Any]] = field(default_factory=list)
    """Extracted tables (empty when ``extract_tables`` is disabled)."""

    detected_languages: dict[str, float] = field(default_factory=dict)
    """Language → confidence mapping (empty when ``detect_language`` is disabled)."""

    page_count: Optional[int] = None
    """Number of pages in the source document (if available)."""


# ---------------------------------------------------------------------------
# Page-level extraction output
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class PageContent:
    """Single page extraction result.

    Returned by :meth:`DocumentExtractor.extract_pages` to represent the
    text content of one PDF page.
    """

    page_number: int
    """1-indexed page number."""

    content: str
    """Extracted text content (may be empty string)."""


# ---------------------------------------------------------------------------
# Document extractor facade
# ---------------------------------------------------------------------------

class DocumentExtractor:
    """Unified document extraction facade over kreuzberg.

    Provides a clean, configurable interface for document text extraction,
    centralizing all kreuzberg interaction within a single module.

    Usage::

        # Basic extraction (identical to legacy fast_extract)
        result = await DocumentExtractor.extract(path)

        # Enhanced extraction with tables and metadata
        result = await DocumentExtractor.extract(path, DocumentExtractor.ENHANCED)

        # Custom profile
        profile = ExtractionProfile(output_format="markdown", extract_tables=True)
        result = await DocumentExtractor.extract(path, profile)
    """

    # Pre-defined profiles -------------------------------------------------

    BASIC: ClassVar[ExtractionProfile] = ExtractionProfile()
    """Plain-text extraction only — equivalent to legacy ``fast_extract()``."""

    ENHANCED: ClassVar[ExtractionProfile] = ExtractionProfile(
        output_format="markdown",
        extract_tables=True,
        extract_metadata=True,
        pdf_extract_metadata=True,
        force_ocr=False,
    )
    """Rich extraction with tables, metadata, and layout-based table detection.

    ``force_ocr`` is disabled because:
    - Most documents (e.g. 10-K, 10-Q PDFs) already contain a native text layer.
    - kreuzberg automatically falls back to OCR for scanned / image-only pages.
    - Forcing OCR triggers Tesseract ObjectCache leak warnings in concurrent use
      and significantly slows down compilation with no quality benefit.
    """

    # Public API -----------------------------------------------------------

    @staticmethod
    async def extract(
        file_path: Union[str, Path],
        profile: Optional[ExtractionProfile] = None,
    ) -> ExtractionOutput:
        """Extract content from a single file.

        Args:
            file_path: Path to the document.
            profile:   Extraction profile.  Defaults to :attr:`BASIC`.

        Returns:
            :class:`ExtractionOutput` with at least ``content`` populated.

        Raises:
            FileNotFoundError: If *file_path* does not exist.
            Exception: Propagates kreuzberg extraction errors after logging.
        """
        from kreuzberg import extract_file

        profile = profile or DocumentExtractor.BASIC
        config = DocumentExtractor._build_config(profile)

        try:
            result = await extract_file(str(file_path), config=config)
            output = DocumentExtractor._convert_result(result, profile)
            # Fallback: kreuzberg 4.9.1 returns page_count=0 when force_ocr=True;
            # use pypdf to get the real page count when missing.
            if output.page_count is None:
                fallback = DocumentExtractor._fallback_page_count(file_path)
                if fallback is not None:
                    output = ExtractionOutput(
                        content=output.content,
                        mime_type=output.mime_type,
                        metadata=output.metadata,
                        tables=output.tables,
                        detected_languages=output.detected_languages,
                        page_count=fallback,
                    )
            return output
        except Exception as exc:
            fallback = DocumentExtractor._fallback_docx(file_path, profile)
            if fallback is not None:
                logger.warning(
                    "Document extraction for {} fell back to python-docx/xml: {}",
                    file_path,
                    exc,
                )
                return fallback
            logger.error(
                "Document extraction failed for {}: {}",
                file_path,
                exc,
            )
            raise

    @staticmethod
    async def extract_isolated(
        file_path: Union[str, Path],
        profile: Optional[ExtractionProfile] = None,
    ) -> ExtractionOutput:
        """Extract content in a fully isolated child process.

        Each call spawns a fresh ``multiprocessing.Process``.  When the
        child exits (normally or via crash), the OS reclaims **all** of
        its native memory — Rust arenas, layout-model buffers, image
        caches — guaranteeing zero accumulation in the parent.

        Unlike ``ProcessPoolExecutor``, a crash in one extraction never
        poisons future calls.

        Falls back to in-process extraction on subprocess failure.
        """
        profile = profile or DocumentExtractor.BASIC
        profile_dict = {
            f.name: getattr(profile, f.name)
            for f in dataclasses.fields(profile)
        }

        loop = asyncio.get_event_loop()
        try:
            raw = await loop.run_in_executor(
                None,
                _run_extraction_in_child,
                str(file_path),
                profile_dict,
            )
            return ExtractionOutput(
                content=raw["content"],
                mime_type=raw.get("mime_type", ""),
                metadata=raw.get("metadata", {}),
                tables=raw.get("tables", []),
                detected_languages=raw.get("detected_languages", {}),
                page_count=raw.get("page_count"),
            )
        except Exception as exc:
            logger.warning(
                "Subprocess extraction failed for {}, falling back to in-process: {}",
                file_path, exc,
            )
            fallback = DocumentExtractor._fallback_docx(file_path, profile)
            if fallback is not None:
                return fallback
            return await DocumentExtractor.extract(file_path, profile)

    @staticmethod
    async def extract_bytes(
        data: bytes,
        mime_type: str,
        profile: Optional[ExtractionProfile] = None,
    ) -> ExtractionOutput:
        """Extract content from raw bytes.

        Args:
            data:      File content as bytes.
            mime_type: MIME type of the data (required for format detection).
            profile:   Extraction profile.  Defaults to :attr:`BASIC`.

        Returns:
            :class:`ExtractionOutput`.
        """
        from kreuzberg import extract_bytes as _extract_bytes

        profile = profile or DocumentExtractor.BASIC
        config = DocumentExtractor._build_config(profile)

        try:
            result = await _extract_bytes(data, mime_type=mime_type, config=config)
            return DocumentExtractor._convert_result(result, profile)
        except Exception:
            logger.error("Byte extraction failed for mime_type={}", mime_type)
            raise

    @staticmethod
    async def batch_extract(
        file_paths: Sequence[Union[str, Path]],
        profile: Optional[ExtractionProfile] = None,
    ) -> List[ExtractionOutput]:
        """Extract content from multiple files in parallel.

        Args:
            file_paths: Sequence of document paths.
            profile:    Extraction profile.  Defaults to :attr:`BASIC`.

        Returns:
            List of :class:`ExtractionOutput`, one per input path.
        """
        from kreuzberg import BatchFileItem, batch_extract_files

        profile = profile or DocumentExtractor.BASIC
        config = DocumentExtractor._build_config(profile)

        try:
            items = [BatchFileItem(str(path), config=config) for path in file_paths]
            results = await batch_extract_files(items, config=config)
            outputs = [
                DocumentExtractor._convert_result(r, profile) for r in results
            ]
            # Apply page_count fallback for each output
            fixed: List[ExtractionOutput] = []
            for output, fp in zip(outputs, file_paths):
                if output.page_count is None:
                    fallback = DocumentExtractor._fallback_page_count(fp)
                    if fallback is not None:
                        output = ExtractionOutput(
                            content=output.content,
                            mime_type=output.mime_type,
                            metadata=output.metadata,
                            tables=output.tables,
                            detected_languages=output.detected_languages,
                            page_count=fallback,
                        )
                fixed.append(output)
            return fixed
        except Exception:
            logger.error("Batch extraction failed for {} files", len(file_paths))
            raise

    # Page-level extraction -------------------------------------------------

    @staticmethod
    def extract_pages(
        file_path: Union[str, Path],
        pages: list[int],
    ) -> list[PageContent]:
        """Extract text content from specific PDF pages.

        Uses pypdf to read individual pages by 1-indexed page number.
        Invalid page numbers (< 1 or > total pages) are silently skipped.

        Args:
            file_path: Path to a PDF file.
            pages:     List of 1-indexed page numbers to extract.

        Returns:
            List of :class:`PageContent` for each valid requested page,
            in the order given by *pages*.

        Raises:
            FileNotFoundError: If *file_path* does not exist.
            Exception: On PDF parsing failure (logged before re-raise).
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"PDF file not found: {path}")

        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            total = len(reader.pages)
            valid_pages = [p for p in pages if 1 <= p <= total]
            return [
                PageContent(
                    page_number=p,
                    content=reader.pages[p - 1].extract_text() or "",
                )
                for p in valid_pages
            ]
        except FileNotFoundError:
            raise
        except Exception as exc:
            logger.error(
                "Page-level extraction failed for {}: {}",
                file_path,
                exc,
            )
            raise

    @staticmethod
    def extract_page_range(
        file_path: Union[str, Path],
        start_page: int,
        end_page: int,
    ) -> list[PageContent]:
        """Extract text content from a contiguous range of PDF pages.

        Convenience wrapper around :meth:`extract_pages`.

        Args:
            file_path:  Path to a PDF file.
            start_page: First page (1-indexed, inclusive).
            end_page:   Last page (1-indexed, inclusive).

        Returns:
            List of :class:`PageContent` for the requested range.
        """
        pages = list(range(start_page, end_page + 1))
        return DocumentExtractor.extract_pages(file_path, pages)

    # Internal helpers -----------------------------------------------------

    @staticmethod
    def _output_format(enum_cls: Any, name: str) -> Any:
        """Return a kreuzberg OutputFormat value across API versions."""
        try:
            return enum_cls(name.lower())
        except Exception:
            return getattr(enum_cls, name.upper(), None)

    @staticmethod
    def _fallback_page_count(
        file_path: Union[str, Path],
    ) -> Optional[int]:
        """Get page count via pypdf when kreuzberg fails to report it.

        kreuzberg >= 4.9.1 returns ``get_page_count() == 0`` when
        ``force_ocr=True`` is set.  This fallback uses pypdf (already a
        transitive dependency) for a lightweight page-count-only read.

        Returns:
            Page count, or None for non-PDF files or on error.
        """
        if Path(file_path).suffix.lower() != ".pdf":
            return None
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(file_path))
            count = len(reader.pages)
            return count if count > 0 else None
        except Exception:
            return None

    @staticmethod
    def _fallback_docx(
        file_path: Union[str, Path],
        profile: ExtractionProfile,
    ) -> Optional[ExtractionOutput]:
        """Extract docx text without kreuzberg's native extension."""
        path = Path(file_path)
        if path.suffix.lower() != ".docx":
            return None

        try:
            from docx import Document

            document = Document(str(path))
            lines: list[str] = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    lines.append(text)

            tables: list[dict[str, Any]] = []
            for table_index, table in enumerate(document.tables):
                table_rows: list[list[str]] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        table_rows.append(cells)
                        lines.append("\t".join(cells))
                if profile.extract_tables and table_rows:
                    tables.append({
                        "page_number": None,
                        "table_index": table_index,
                        "rows": table_rows,
                    })

            content = "\n".join(lines).strip()
            if not content:
                return None

            metadata: dict[str, Any] = {}
            if profile.extract_metadata:
                props = document.core_properties
                metadata = {
                    key: value
                    for key, value in {
                        "author": props.author,
                        "title": props.title,
                        "subject": props.subject,
                        "keywords": props.keywords,
                        "created": props.created.isoformat() if props.created else None,
                        "modified": props.modified.isoformat() if props.modified else None,
                    }.items()
                    if value
                }

            return ExtractionOutput(
                content=content,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                metadata=metadata,
                tables=tables,
            )
        except Exception as exc:
            logger.debug("python-docx fallback failed for {}: {}", path, exc)

        try:
            with zipfile.ZipFile(path) as archive:
                xml_data = archive.read("word/document.xml")
            root = ET.fromstring(xml_data)
            namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            lines = []
            for paragraph in root.findall(".//w:p", namespace):
                text = "".join(
                    node.text or "" for node in paragraph.findall(".//w:t", namespace)
                ).strip()
                if text:
                    lines.append(text)
            content = "\n".join(lines).strip()
            if not content:
                return None
            return ExtractionOutput(
                content=content,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as exc:
            logger.debug("docx xml fallback failed for {}: {}", path, exc)
            return None

    @staticmethod
    def _build_config(profile: ExtractionProfile):
        """Build a kreuzberg ``ExtractionConfig`` from an :class:`ExtractionProfile`.

        Maps profile fields to the kreuzberg configuration objects that are
        actually available in the installed version.
        """
        from kreuzberg import (
            ExtractionConfig,
            OcrConfig,
            OutputFormat,
            PageConfig,
            PdfConfig,
        )

        # --- Output format ---
        format_map = {
            "plain": DocumentExtractor._output_format(OutputFormat, "plain"),
            "markdown": DocumentExtractor._output_format(OutputFormat, "markdown"),
            "html": DocumentExtractor._output_format(OutputFormat, "html"),
            "djot": DocumentExtractor._output_format(OutputFormat, "djot"),
        }
        output_format = format_map.get(profile.output_format) or format_map["plain"]

        # --- OCR config ---
        ocr_config: Optional[OcrConfig] = None
        if profile.ocr_enabled:
            ocr_config = OcrConfig(
                backend=profile.ocr_backend,
                language=profile.ocr_language,
            )

        # --- Page config ---
        page_config: Optional[PageConfig] = None
        if profile.page_start is not None or profile.page_end is not None:
            # kreuzberg PageConfig.extract_pages expects a list of page indices
            pages: Optional[list[int]] = None
            if profile.page_start is not None:
                end = profile.page_end if profile.page_end is not None else profile.page_start
                pages = list(range(profile.page_start, end + 1))
            page_config = PageConfig(extract_pages=pages)

        # --- PDF config ---
        pdf_config: Optional[PdfConfig] = None
        if any([
            profile.pdf_extract_images,
            profile.pdf_extract_metadata,
            profile.pdf_password,
        ]):
            passwords = [profile.pdf_password] if profile.pdf_password else None
            pdf_config = PdfConfig(
                extract_images=profile.pdf_extract_images,
                extract_metadata=profile.pdf_extract_metadata,
                passwords=passwords,
            )

        # --- Language detection ---
        lang_config = None
        if profile.detect_language:
            from kreuzberg import LanguageDetectionConfig
            lang_config = LanguageDetectionConfig(enabled=True)

        # --- Layout detection for table extraction ---
        layout_config = None
        if profile.extract_tables:
            try:
                from kreuzberg import LayoutDetectionConfig
                layout_config = LayoutDetectionConfig(
                    confidence_threshold=0.3,
                    apply_heuristics=True,
                    table_model="slanet_auto",
                )
            except ImportError:
                pass

        # --- Assemble ExtractionConfig ---
        kwargs: dict[str, Any] = {
            "output_format": output_format,
        }
        if ocr_config is not None:
            kwargs["ocr"] = ocr_config
        if profile.force_ocr:
            kwargs["force_ocr"] = True
        elif profile.force_ocr_pages:
            kwargs["force_ocr_pages"] = list(profile.force_ocr_pages)
        if page_config is not None:
            kwargs["pages"] = page_config
        if pdf_config is not None:
            kwargs["pdf_options"] = pdf_config
        if lang_config is not None:
            kwargs["language_detection"] = lang_config
        if profile.max_concurrent is not None:
            kwargs["max_concurrent_extractions"] = profile.max_concurrent
        if layout_config is not None:
            kwargs["layout"] = layout_config

        return ExtractionConfig(**kwargs)

    @staticmethod
    def _convert_result(
        result: "ExtractionResult",
        profile: ExtractionProfile,
    ) -> ExtractionOutput:
        """Convert a kreuzberg ``ExtractionResult`` to :class:`ExtractionOutput`.

        Only populates optional fields when the corresponding profile flag is
        enabled, keeping the output lean for basic extraction.
        """
        content: str = result.content or ""
        mime_type: str = getattr(result, "mime_type", "") or ""

        # Metadata
        metadata: dict[str, Any] = {}
        if profile.extract_metadata:
            raw_meta = getattr(result, "metadata", None)
            if raw_meta is not None:
                if isinstance(raw_meta, dict):
                    metadata = dict(raw_meta)
                else:
                    # kreuzberg may return a non-dict metadata object
                    try:
                        metadata = dict(raw_meta)
                    except (TypeError, ValueError):
                        metadata = {"raw": str(raw_meta)}

        # Tables
        tables: list[dict[str, Any]] = []
        if profile.extract_tables:
            raw_tables = getattr(result, "tables", None) or []
            for t in raw_tables:
                if isinstance(t, dict):
                    tables.append(t)
                else:
                    # kreuzberg ExtractedTable has: cells, markdown, page_number
                    tables.append({
                        "markdown": getattr(t, "markdown", ""),
                        "cells": getattr(t, "cells", []),
                        "page_number": getattr(t, "page_number", None),
                    })

        # Language detection
        detected_languages: dict[str, float] = {}
        if profile.detect_language:
            raw_langs = getattr(result, "detected_languages", None)
            if raw_langs:
                for entry in raw_langs:
                    if isinstance(entry, dict):
                        lang = entry.get("language", "")
                        conf = entry.get("confidence", 0.0)
                    else:
                        # kreuzberg DetectedLanguage object
                        lang = getattr(entry, "language", "")
                        conf = getattr(entry, "confidence", 0.0)
                    if lang:
                        detected_languages[lang] = float(conf)

        # Page count — prefer get_page_count() over get_chunk_count()
        page_count: Optional[int] = None
        get_page_count = getattr(result, "get_page_count", None)
        if get_page_count and callable(get_page_count):
            cnt = get_page_count()
            if cnt is not None and cnt > 0:
                page_count = cnt

        return ExtractionOutput(
            content=content,
            mime_type=mime_type,
            metadata=metadata,
            tables=tables,
            detected_languages=detected_languages,
            page_count=page_count,
        )
